from rest_framework import viewsets, status, permissions
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from django.http import HttpResponse
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils import timezone
from io import BytesIO
from docx import Document
from docx.shared import Pt
import os
import re

from .models import DocumentTemplate, GeneratedDocument
from .serializers import (
	DocumentTemplateSerializer, 
	GeneratedDocumentSerializer,
	ResumeGenerationSerializer
)


class DocumentTemplateViewSet(viewsets.ReadOnlyModelViewSet):
	"""Read-only viewset for available document templates."""
	queryset = DocumentTemplate.objects.filter(is_active=True)
	serializer_class = DocumentTemplateSerializer
	permission_classes = [IsAuthenticated]


class GeneratedDocumentViewSet(viewsets.ModelViewSet):
	"""CRUD for user-generated documents. Users can only access their own documents."""
	serializer_class = GeneratedDocumentSerializer
	permission_classes = [IsAuthenticated]

	def get_queryset(self):
		user = self.request.user
		return GeneratedDocument.objects.filter(user=user).order_by("-created_at")

	def create(self, request, *args, **kwargs):
		serializer = self.get_serializer(data=request.data)
		serializer.is_valid(raise_exception=True)
		# extract file if provided
		upload_file = serializer.validated_data.pop("file", None)
		instance = serializer.save(user=request.user)

		if upload_file:
			upload_path = os.path.join("documents", str(request.user.id), upload_file.name)
			saved_path = default_storage.save(upload_path, upload_file)
			# build accessible URL (MEDIA_URL + saved_path)
			file_url = settings.MEDIA_URL.rstrip("/") + "/" + saved_path.replace("\\", "/")
			instance.file_url = file_url
			instance.save()

		out_serializer = self.get_serializer(instance)
		return Response(out_serializer.data, status=status.HTTP_201_CREATED)

	def update(self, request, *args, **kwargs):
		partial = kwargs.pop("partial", False)
		instance = self.get_object()
		serializer = self.get_serializer(instance, data=request.data, partial=partial)
		serializer.is_valid(raise_exception=True)
		upload_file = serializer.validated_data.pop("file", None)
		instance = serializer.save()

		if upload_file:
			upload_path = os.path.join("documents", str(request.user.id), upload_file.name)
			saved_path = default_storage.save(upload_path, upload_file)
			file_url = settings.MEDIA_URL.rstrip("/") + "/" + saved_path.replace("\\", "/")
			instance.file_url = file_url
			instance.save()

		out_serializer = self.get_serializer(instance)
		return Response(out_serializer.data)

# ===================================================================
# 이력서 자동 생성 API View
# ===================================================================

class ResumeGenerateView(APIView):
	"""
	이력서 자동 생성 API
	- 사용자가 입력한 정보를 바탕으로 docx 형식의 이력서를 생성
	- save_to_documents=true 시 GeneratedDocument로 저장
	"""
	permission_classes = [IsAuthenticated]

	def post(self, request, *args, **kwargs):
		serializer = ResumeGenerationSerializer(data=request.data)
		serializer.is_valid(raise_exception=True)
		data = serializer.validated_data

		buffer = self._build_docx(data)
		filename = self._build_filename(data.get("document_title", "이력서"))

		saved_file_url = None
		if data.get("save_to_documents"):
			path = os.path.join("documents", str(request.user.id), "resumes", filename)
			saved_path = default_storage.save(path, ContentFile(buffer.getvalue()))
			saved_file_url = settings.MEDIA_URL.rstrip("/") + "/" + saved_path.replace("\\", "/")
			GeneratedDocument.objects.create(
				user=request.user,
				doc_type="이력서",
				title=data.get("document_title") or filename,
				status=data.get("status") or "완료",
				file_url=saved_file_url,
				filled_data_json={k: v for k, v in data.items() if k not in ["save_to_documents", "document_title", "status"]},
			)

		response = HttpResponse(
			buffer.getvalue(),
			content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)
		response["Content-Disposition"] = f'attachment; filename="{filename}"'
		if saved_file_url:
			response["X-Document-Saved"] = "true"
			response["X-Document-Url"] = saved_file_url
		return response

	def _build_filename(self, title: str) -> str:
		"""안전한 파일명 생성"""
		safe = re.sub(r"[^A-Za-z0-9_-]+", "_", title or "resume").strip("_")
		timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
		return f"{safe or 'resume'}_{timestamp}.docx"

	def _build_docx(self, data) -> BytesIO:
		"""이력서 docx 파일 생성"""
		doc = Document()

		def add_title(text, size=18, bold=True):
			p = doc.add_paragraph()
			r = p.add_run(text)
			r.bold = bold
			r.font.size = Pt(size)

		def add_label_value(label: str, value: str):
			if not value:
				return
			p = doc.add_paragraph()
			r = p.add_run(f"{label}: {value}")
			r.bold = False
			r.font.size = Pt(11)

		def add_section(title: str):
			p = doc.add_paragraph()
			r = p.add_run(title)
			r.bold = True
			r.font.size = Pt(14)

		def add_bullet(text: str):
			if text:
				p = doc.add_paragraph(text, style="List Bullet")
				p.runs[0].font.size = Pt(11)

		# Header
		add_title(data.get("name", ""), size=20)
		add_label_value("직무", data.get("title", ""))
		add_label_value("연락처", data.get("phone", ""))
		add_label_value("이메일", data.get("email", ""))
		add_label_value("주소", data.get("address", ""))

		if data.get("summary"):
			add_section("요약")
			doc.add_paragraph(data.get("summary", ""))

		if data.get("experiences"):
			add_section("경력")
			for exp in data.get("experiences", []):
				company_line = " ".join([exp.get("company", ""), exp.get("role", "")]).strip()
				period = exp.get("period", "")
				title_line = company_line if not period else f"{company_line} ({period})"
				if title_line:
					doc.add_paragraph(title_line, style="List Number")
				if exp.get("description"):
					doc.add_paragraph(exp.get("description"))
				for a in exp.get("achievements", []) or []:
					add_bullet(a)

		if data.get("educations"):
			add_section("학력")
			for edu in data.get("educations", []):
				school_line = edu.get("school", "")
				period = edu.get("period", "")
				if period:
					school_line = f"{school_line} ({period})"
				if school_line:
					doc.add_paragraph(school_line, style="List Number")
				if edu.get("major"):
					add_bullet(f"전공: {edu.get('major')}")
				if edu.get("description"):
					add_bullet(edu.get("description"))

		if data.get("skills"):
			add_section("기술")
			add_label_value("보유 기술", ", ".join([s for s in data.get("skills", []) if s]))

		if data.get("certifications"):
			add_section("자격증/수상")
			for cert in data.get("certifications", []):
				add_bullet(cert)

		if data.get("languages"):
			add_section("언어")
			add_label_value("언어", ", ".join([l for l in data.get("languages", []) if l]))

		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		return buffer
